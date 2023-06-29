from docx import Document
from docx.shared import Pt


def generate_quotation_document(main_request, room_requests):
    # Create a new Word document
    document = Document()

    # Add a title
    document.add_heading('Quotation', level=1)

    # Add main request details
    document.add_heading('Main Request', level=2)
    document.add_paragraph(f'Country: {main_request.country}')
    document.add_paragraph(f'Area: {main_request.area}')

    # Add room requests details
    document.add_heading('Room Requests', level=2)
    for i, room_request in enumerate(room_requests, start=1):
        document.add_paragraph(f'Room {i}')
        document.add_paragraph(f'Check-in: {room_request.check_in}')
        document.add_paragraph(f'Check-out: {room_request.check_out}')
        document.add_paragraph(f'Number of Adults: {room_request.num_adults}')
        document.add_paragraph(f'Number of Children: {room_request.num_children}')
        # Add more room request details as needed

    # Save the document
    document.save('quotation.docx')


